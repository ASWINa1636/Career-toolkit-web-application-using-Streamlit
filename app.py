import streamlit as st
from docx import Document
from docx.shared import Pt, Inches
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
import fitz  # PyMuPDF
import re
import io

# ---------------------- HELPER FUNCTIONS ---------------------- #

def extract_text_from_file(uploaded_file):
    text = ""
    try:
        if uploaded_file.type == "application/pdf":
            with fitz.open(stream=uploaded_file.read(), filetype="pdf") as doc:
                for page in doc:
                    text += page.get_text()
        elif uploaded_file.type == "application/vnd.openxmlformats-officedocument.wordprocessingml.document":
            doc = Document(uploaded_file)
            text = "\n".join([p.text for p in doc.paragraphs])
    except Exception as e:
        st.error(f"Error reading file: {e}")
    return text

def score_resume(text):
    text = text.lower()
    score = 0
    
    # --- Skills Section (Max 30) ---
    skill_keywords = ["Python", "Java", "JavaScript", "C++", "C#", "Ruby", "Go", "Swift", "Kotlin", "TypeScript", "PHP", "Rust",
    "Scala", "Perl", "Haskell", "Lua", "Linux", "Windows", "macOS", "UNIX", "MATLAB", "Power Systems", "HTML", 
    "CSS", "React", "Angular", "Vue.js", "Node.js", "Django", "Flask", "Spring Boot", "ASP.NET", "Laravel",
    "Machine Learning", "Deep Learning", "Data Science", "TensorFlow", "PyTorch", "Keras", "Pandas", "NumPy", 
    "Scikit-Learn", "R", "Matplotlib", "Seaborn", "OpenAI API", "Natural Language Processing", "Computer Vision",
    "Big Data", "SQL", "PostgreSQL", "MongoDB", "Firebase", "AWS", "Azure", "Google Cloud", "Docker", "Kubernetes",
    "Ethical Hacking", "Penetration Testing", "Cryptography", "Network Security", "SOC Analyst", "Malware Analysis",
    "Reverse Engineering", "CI/CD", "Jenkins", "Terraform", "Ansible", "Git", "GitHub Actions", "GitLab CI", 
    "Bash Scripting", "PowerShell", "Agile", "Scrum", "Kanban", "JIRA", "Trello", "Confluence", "PCB Design", 
    "Web Development", "Mobile Development", "Word", "Excel", "PowerPoint", "Outlook", "Tableau", "Power BI", 
    "Apache Spark", "Hadoop", "Kafka", "Elasticsearch", "GraphQL", "REST APIs", "SOAP", "Microservices", 
    "DevOps", "System Administration", "Virtualization", "VMware", "Hyper-V", "Cloud Security", "IoT", 
    "Embedded Systems", "Arduino", "Raspberry Pi", "Blockchain", "Solidity", "UI/UX Design", "Figma", "Adobe XD", 
    "Photoshop", "Illustrator", "Blender", "3D Modeling", "Game Development", "Unity", "Unreal Engine", "OpenGL", 
    "WebAssembly", "Quantum Computing", "Statistics", "Probability", "Linear Algebra", "Data Visualization", 
    "ETL Processes", "Data Warehousing", "Snowflake", "Redshift", "DynamoDB", "Cassandra", "Neo4j", "Redis", 
    "Load Balancing", "NGINX", "Apache", "Incident Response", "Forensic Analysis", "Threat Hunting", 
    "Cybersecurity Frameworks", "NIST", "ISO 27001", "GDPR Compliance", "Project Management", "PMP", "Lean Six Sigma", 
    "Technical Writing", "Public Speaking", "Team Leadership", "Conflict Resolution", "Time Management", 
    "Customer Relationship Management (CRM)", "Salesforce", "SAP", "ERP Systems", "Supply Chain Management", 
    "Digital Marketing", "SEO", "SEM", "Content Management Systems (CMS)", "WordPress", "Shopify", "Magento", 
    "Augmented Reality (AR)", "Virtual Reality (VR)", "Robotics", "ROS (Robot Operating System)", "PLC Programming", 
    "AutoCAD", "SolidWorks", "Finite Element Analysis (FEA)", "Computational Fluid Dynamics (CFD)", "Simulink", 
    "VLSI Design", "Verilog", "VHDL", "FPGA Programming", "Signal Processing", "Image Processing", "Audio Engineering", 
    "Penetration Testing Tools (Metasploit, Burp Suite)", "Wireshark", "Nmap", "Splunk", "SIEM", "Log Analysis", 
    "Chaos Engineering", "Site Reliability Engineering (SRE)", "Monitoring Tools (Prometheus, Grafana)", 
    "Version Control Systems", "Subversion (SVN)", "Mercurial", "Test Automation", "Selenium", "Cypress", 
    "Postman", "Unit Testing", "Integration Testing", "Performance Testing", "Load Testing", "Stress Testing", 
    "Behavior-Driven Development (BDD)", "Test-Driven Development (TDD)", "Pair Programming", "Code Review", 
    "Documentation", "API Design", "OAuth", "JWT", "Microfrontend", "Serverless Architecture", "Edge Computing", 
    "Bioinformatics", "Genomics", "Proteomics", "Molecular Modeling", "Chemoinformatics", "Financial Modeling", 
    "Risk Analysis", "Algorithm Design", "Data Structures", "Competitive Programming", "Parallel Computing", 
    "Distributed Systems", "Graph Theory", "Optimization", "Simulation", "Forecasting", "Econometrics", 
    "Geospatial Analysis", "GIS (Geographic Information Systems)", "Remote Sensing", "Satellite Imagery Analysis", 
    "Drone Technology", "Aeronautical Engineering", "Mechanical Design", "Thermodynamics", "Materials Science", 
    "Nanotechnology", "Renewable Energy Systems", "Solar Technology", "Wind Energy", "Battery Systems", 
    "Electrical Engineering", "Control Systems", "Power Electronics", "RF Engineering", "Antenna Design", 
    "Satellite Communications", "5G Technology", "Network Protocols", "TCP/IP", "DNS Management", "VPN Configuration", 
    "Customer Support", "Technical Support", "ITIL", "ServiceNow", "Help Desk Management", "Change Management", 
    "Disaster Recovery", "Business Continuity Planning", "Stakeholder Management", "Negotiation", "Critical Thinking", 
    "Problem Solving", "Emotional Intelligence", "Adaptability", "Cross-Functional Collaboration", "Mentoring", 
    "Training & Development", "Instructional Design", "E-Learning Development", "LMS (Learning Management Systems)"]
    # Simple count mechanism
    found_skills = [skill for skill in skill_keywords if skill in text]
    skill_count = len(set(found_skills)) # Unique skills
    score += min(skill_count * 2, 30)

    # --- Education (Max 20) ---
    if re.search(r'\b(b\.tech|btech|bachelor|be|mtech|m\.tech|msc|mca|bsc|degree)\b', text):
        score += 15
    if 'cgpa' in text or 'percentage' in text:
        score += 5

    # --- Experience (Max 20) ---
    if 'intern' in text or 'experience' in text or 'project' in text:
        score += 15
    if 'company' in text or 'organization' in text:
        score += 5

    # --- Contact Info (Max 15) ---
    if re.search(r'[\w\.-]+@[\w\.-]+', text):  # Email
        score += 8
    if re.search(r'\b\d{10}\b', text):  # Phone number
        score += 7

    # --- Structure (Max 15) ---
    sections = ['skills', 'education', 'experience']
    found_sections = sum(1 for s in sections if s in text)
    if found_sections == 3:
        score += 10
    elif found_sections > 0:
        score += 5

    return min(score, 100), found_skills

# Function to add hyperlink to DOCX
def add_hyperlink(paragraph, url, text):
    part = paragraph.part
    r_id = part.relate_to(url, "http://schemas.openxmlformats.org/officeDocument/2006/relationships/hyperlink", is_external=True)
    hyperlink = OxmlElement("w:hyperlink")
    hyperlink.set(qn("r:id"), r_id)
    new_run = OxmlElement("w:r")
    rPr = OxmlElement("w:rPr")
    color = OxmlElement("w:color")
    color.set(qn("w:val"), "0000FF")
    rPr.append(color)
    underline = OxmlElement("w:u")
    underline.set(qn("w:val"), "single")
    rPr.append(underline)
    new_run.append(rPr)
    t = OxmlElement("w:t")
    t.text = text
    new_run.append(t)
    hyperlink.append(new_run)
    paragraph._p.append(hyperlink)

# ---------------------- STREAMLIT UI ---------------------- #

st.set_page_config(page_title="Career Toolkit", page_icon="💼", layout="wide")

st.sidebar.title("Career Toolkit 🚀")
menu = st.sidebar.radio("Navigation", ["ATS Scanner", "Resume Creator", "CGPA Calculator"])

# ================= ATS SCANNER =================
if menu == "ATS Scanner":
    st.header("📄 ATS Resume Scanner")
    st.write("Upload your resume to see how well it scores against common keywords.")

    uploaded_file = st.file_uploader("Upload Resume (PDF or DOCX)", type=['pdf', 'docx'])

    if uploaded_file is not None:
        with st.spinner("Scanning..."):
            text = extract_text_from_file(uploaded_file)
            score, skills_found = score_resume(text)
            
            st.divider()
            col1, col2 = st.columns([1, 2])
            
            with col1:
                st.metric(label="ATS Score", value=f"{score}/100")
                if score >= 80:
                    st.success("Excellent Score!")
                elif score >= 50:
                    st.warning("Good, but needs improvement.")
                else:
                    st.error("Needs significant improvement.")
            
            with col2:
                st.subheader("Skills Detected")
                if skills_found:
                    st.write(", ".join(skills_found).title())
                else:
                    st.write("No specific technical keywords detected.")

# ================= RESUME CREATOR =================
elif menu == "Resume Creator":
    st.header("📝 Resume Creator")
    
    # Initialize session state lists if they don't exist
    if 'edu_list' not in st.session_state: st.session_state.edu_list = []
    if 'exp_list' not in st.session_state: st.session_state.exp_list = []
    if 'proj_list' not in st.session_state: st.session_state.proj_list = []
    if 'cert_list' not in st.session_state: st.session_state.cert_list = []
    if 'custom_skills_list' not in st.session_state: st.session_state.custom_skills_list = []

    # --- Tabs for organization ---
    tab1, tab2, tab3, tab4, tab5, tab6 = st.tabs(["Personal Info", "Education", "Experience", "Projects", "Skills & Certs", "Download"])

    with tab1:
        st.subheader("Personal Information")
        col1, col2 = st.columns(2)
        name = col1.text_input("Full Name")
        phone = col2.text_input("Phone Number")
        email = col1.text_input("Email")
        linkedin = col2.text_input("LinkedIn URL")
        github = col1.text_input("GitHub URL")
        portfolio = col2.text_input("Portfolio URL")
        about = st.text_area("About Me / Summary")

    with tab2:
        st.subheader("Education")
        with st.form("edu_form", clear_on_submit=True):
            e_deg = st.text_input("Degree/School")
            e_grade = st.text_input("CGPA/Percentage (Optional)")
            e_year = st.text_input("Year (e.g., 2020-2024)")
            submitted_edu = st.form_submit_button("Add Education")
            if submitted_edu and e_deg:
                st.session_state.edu_list.append({"degree": e_deg, "grade": e_grade, "year": e_year})
        
        # Display added items
        if st.session_state.edu_list:
            st.write("Added Education:")
            for i, item in enumerate(st.session_state.edu_list):
                st.info(f"{i+1}. {item['degree']} | {item['year']}")
                if st.button(f"Remove Edu {i+1}", key=f"del_edu_{i}"):
                    st.session_state.edu_list.pop(i)
                    st.rerun()

    with tab3:
        st.subheader("Experience")
        with st.form("exp_form", clear_on_submit=True):
            exp_role = st.text_input("Role")
            exp_comp = st.text_input("Company")
            exp_dur = st.text_input("Duration")
            exp_desc = st.text_area("Description")
            submitted_exp = st.form_submit_button("Add Experience")
            if submitted_exp and exp_role:
                st.session_state.exp_list.append({"role": exp_role, "company": exp_comp, "duration": exp_dur, "desc": exp_desc})

        if st.session_state.exp_list:
            st.write("Added Experience:")
            for i, item in enumerate(st.session_state.exp_list):
                st.info(f"{i+1}. {item['role']} at {item['company']}")
                if st.button(f"Remove Exp {i+1}", key=f"del_exp_{i}"):
                    st.session_state.exp_list.pop(i)
                    st.rerun()

    with tab4:
        st.subheader("Projects")
        with st.form("proj_form", clear_on_submit=True):
            p_title = st.text_input("Project Title")
            p_tech = st.text_input("Tech Stack")
            p_desc = st.text_area("Description")
            submitted_proj = st.form_submit_button("Add Project")
            if submitted_proj and p_title:
                st.session_state.proj_list.append({"title": p_title, "tech": p_tech, "desc": p_desc})

        if st.session_state.proj_list:
            st.write("Added Projects:")
            for i, item in enumerate(st.session_state.proj_list):
                st.info(f"{i+1}. {item['title']}")
                if st.button(f"Remove Proj {i+1}", key=f"del_proj_{i}"):
                    st.session_state.proj_list.pop(i)
                    st.rerun()

    with tab5:
        st.subheader("Skills & Certifications")
        lang = st.text_input("Languages (e.g., Python, Java)")
        tools = st.text_input("Developer Tools")
        libs = st.text_input("Libraries/Frameworks")
        soft = st.text_input("Soft Skills")
        
        st.markdown("---")
        st.write("Add Certifications")
        c_input = st.text_input("Certification Name")
        if st.button("Add Certification"):
            if c_input:
                st.session_state.cert_list.append(c_input)
        
        if st.session_state.cert_list:
             st.write(st.session_state.cert_list)

    with tab6:
        st.subheader("Generate Resume")
        filename = st.text_input("Filename", value="MyResume")
        
        if st.button("Generate DOCX"):
            # Create DOCX
            doc = Document()
            section = doc.sections[0]
            section.top_margin = Inches(0.5)
            section.bottom_margin = Inches(0.5)
            section.left_margin = Inches(0.5)
            section.right_margin = Inches(0.5)

            # Name
            name_para = doc.add_paragraph(name)
            name_para.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
            name_para.runs[0].font.size = Pt(20)
            name_para.runs[0].bold = True
            
            # Contact
            contact_para = doc.add_paragraph()
            contact_para.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
            contact_para.add_run(f"{phone} | {email}")
            
            if linkedin:
                contact_para.add_run(" | ")
                add_hyperlink(contact_para, linkedin, "LinkedIn")
            if github:
                contact_para.add_run(" | ")
                add_hyperlink(contact_para, github, "GitHub")
            if portfolio:
                contact_para.add_run(" | ")
                add_hyperlink(contact_para, portfolio, "Portfolio")

            # About
            if about:
                doc.add_heading("About Me", level=1)
                doc.add_paragraph(about)

            # Education
            if st.session_state.edu_list:
                doc.add_heading("Education", level=1)
                for item in st.session_state.edu_list:
                    p = doc.add_paragraph(style='List Bullet')
                    runner = f"{item['degree']}, {item['year']}"
                    if item['grade']:
                         runner += f" (Grade: {item['grade']})"
                    p.add_run(runner)

            # Experience
            if st.session_state.exp_list:
                doc.add_heading("Experience", level=1)
                for item in st.session_state.exp_list:
                    p = doc.add_paragraph(style='List Bullet')
                    p.add_run(f"{item['role']} at {item['company']} ({item['duration']})").bold = True
                    doc.add_paragraph(item['desc'], style='List Continue')

            # Projects
            if st.session_state.proj_list:
                doc.add_heading("Projects", level=1)
                for item in st.session_state.proj_list:
                    p = doc.add_paragraph(style='List Bullet')
                    p.add_run(f"{item['title']}").bold = True
                    if item['tech']:
                        p.add_run(f" [{item['tech']}]").italic = True
                    doc.add_paragraph(item['desc'], style='List Continue')

            # Skills
            doc.add_heading("Skills", level=1)
            if lang: doc.add_paragraph(f"Languages: {lang}", style='List Bullet')
            if tools: doc.add_paragraph(f"Tools: {tools}", style='List Bullet')
            if libs: doc.add_paragraph(f"Libraries: {libs}", style='List Bullet')
            if soft: doc.add_paragraph(f"Soft Skills: {soft}", style='List Bullet')

            # Certifications
            if st.session_state.cert_list:
                doc.add_heading("Certifications", level=1)
                for cert in st.session_state.cert_list:
                    doc.add_paragraph(cert, style='List Bullet')

            # Save to buffer
            buffer = io.BytesIO()
            doc.save(buffer)
            buffer.seek(0)
            
            st.success("Resume Generated successfully!")
            st.download_button(
                label="📥 Download Resume (.docx)",
                data=buffer,
                file_name=f"{filename}.docx",
                mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document"
            )

# ================= CGPA CALCULATOR =================
elif menu == "CGPA Calculator":
    st.header("🧮 SRM CGPA Calculator")
    
    num_sems = st.number_input("Number of Semesters Completed", min_value=1, max_value=10, value=1, step=1)
    
    st.write("Enter GPA for each semester:")
    
    # Create columns dynamically
    cols = st.columns(4) 
    gpa_values = []
    
    for i in range(1, num_sems + 1):
        # Distribute inputs across columns
        with cols[(i-1) % 4]:
            val = st.number_input(f"Sem {i} GPA", min_value=0.0, max_value=10.0, step=0.01, key=f"gpa_{i}")
            gpa_values.append(val)
            
    if st.button("Calculate CGPA"):
        # Filter out 0.0 values if user forgot to enter them, or keep them if they actually got 0
        # Assuming user enters valid data
        final_cgpa = sum(gpa_values) / len(gpa_values)
        st.divider()
        st.markdown(f"### 🎓 Your CGPA is: `{final_cgpa:.2f}`")